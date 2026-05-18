"""TDD – agents/docs_agent.py"""
import os
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from agents.docs_agent import generate_guest_form, _build_default_doc, _fill_template
from agents.gmail_agent import Reservation


def _make_reservation(**kwargs):
    defaults = dict(
        guest_name="Ana Paula", guest_email="ana@email.com", guest_phone="11987654321",
        checkin="2026-06-01", checkout="2026-06-05",
        room="202", platform="airbnb", booking_id="AB999", is_cancellation=False,
    )
    defaults.update(kwargs)
    return Reservation(**defaults)


class TestBuildDefaultDoc:
    def test_cria_documento_com_cabecalho(self):
        doc = Document()
        res = _make_reservation()
        _build_default_doc(doc, res)
        # add_heading(text, 0) usa estilo "Title"; níveis 1+ usam "Heading N"
        titulo_styles = {"Title", "Heading 1", "Heading 2"}
        headings = [p.text for p in doc.paragraphs if p.style.name in titulo_styles]
        assert any("Ficha" in h for h in headings)

    def test_cria_tabela_com_campos_da_reserva(self):
        doc = Document()
        res = _make_reservation()
        _build_default_doc(doc, res)
        assert len(doc.tables) == 1
        table = doc.tables[0]
        labels = [row.cells[0].text for row in table.rows]
        assert "Nome Completo" in labels
        assert "Check-in" in labels
        assert "Check-out" in labels

    def test_valores_da_reserva_na_tabela(self):
        doc = Document()
        res = _make_reservation()
        _build_default_doc(doc, res)
        table = doc.tables[0]
        values = [row.cells[1].text for row in table.rows]
        assert "Ana Paula" in values
        assert "AB999" in values

    def test_linha_de_assinatura_presente(self):
        doc = Document()
        _build_default_doc(doc, _make_reservation())
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Assinatura" in all_text


class TestFillTemplate:
    def test_substitui_placeholders_nos_paragrafos(self):
        doc = Document()
        doc.add_paragraph("Hóspede: {{guest_name}}")
        doc.add_paragraph("Check-in: {{checkin}}")

        _fill_template(doc, _make_reservation())

        texts = [p.text for p in doc.paragraphs]
        assert any("Ana Paula" in t for t in texts)
        assert any("2026-06-01" in t for t in texts)
        assert not any("{{guest_name}}" in t for t in texts)

    def test_substitui_hotel_name(self):
        doc = Document()
        doc.add_paragraph("Bem-vindo ao {{hotel_name}}!")

        with patch("agents.docs_agent.settings") as mock_settings:
            mock_settings.hotel_name = "Grand Hotel"
            _fill_template(doc, _make_reservation())

        assert any("Grand Hotel" in p.text for p in doc.paragraphs)

    def test_substitui_placeholders_em_celulas_de_tabela(self):
        """Regressão: _fill_template antes só iterava doc.paragraphs, ignorando tabelas."""
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Nome"
        table.rows[0].cells[1].paragraphs[0].runs[0].text if table.rows[0].cells[1].paragraphs[0].runs else None
        # Adiciona run com placeholder na célula
        cell = table.rows[0].cells[1]
        cell.paragraphs[0].clear()
        cell.paragraphs[0].add_run("{{guest_name}}")

        _fill_template(doc, _make_reservation())

        cell_text = table.rows[0].cells[1].paragraphs[0].text
        assert cell_text == "Ana Paula"
        assert "{{guest_name}}" not in cell_text


class TestGenerateGuestForm:
    def test_gera_arquivo_docx(self, tmp_path):
        res = _make_reservation()

        with patch("agents.docs_agent.OUTPUT_DIR", tmp_path), \
             patch("agents.docs_agent.TEMPLATE_PATH", tmp_path / "nao_existe.docx"):
            path = generate_guest_form(res)

        assert Path(path).exists()
        assert path.endswith(".docx")

    def test_nome_do_arquivo_contem_booking_id_e_nome(self, tmp_path):
        res = _make_reservation()

        with patch("agents.docs_agent.OUTPUT_DIR", tmp_path), \
             patch("agents.docs_agent.TEMPLATE_PATH", tmp_path / "nao_existe.docx"):
            path = generate_guest_form(res)

        filename = Path(path).name
        assert "AB999" in filename
        assert "Ana_Paula" in filename

    def test_usa_template_se_existir(self, tmp_path):
        template_path = tmp_path / "ficha_template.docx"
        doc_template = Document()
        doc_template.add_paragraph("Nome: {{guest_name}}")
        doc_template.save(template_path)

        res = _make_reservation()

        with patch("agents.docs_agent.OUTPUT_DIR", tmp_path), \
             patch("agents.docs_agent.TEMPLATE_PATH", template_path):
            path = generate_guest_form(res)

        result_doc = Document(path)
        all_text = " ".join(p.text for p in result_doc.paragraphs)
        assert "Ana Paula" in all_text
        assert "{{guest_name}}" not in all_text
