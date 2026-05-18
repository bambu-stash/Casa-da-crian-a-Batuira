"""
Sub-agente de documentação: gera ficha de hóspede em .docx.
"""
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

from config import settings

OUTPUT_DIR = Path(__file__).parent.parent / "data" / "fichas"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

TEMPLATE_PATH = Path(__file__).parent.parent / "data" / "ficha_template.docx"


def generate_guest_form(reservation) -> str:
    """Gera ficha .docx e retorna o caminho do arquivo gerado."""
    doc = Document(TEMPLATE_PATH) if TEMPLATE_PATH.exists() else Document()

    if not TEMPLATE_PATH.exists():
        _build_default_doc(doc, reservation)
    else:
        _fill_template(doc, reservation)

    filename = f"ficha_{reservation.booking_id}_{reservation.guest_name.replace(' ', '_')}.docx"
    output_path = OUTPUT_DIR / filename
    doc.save(output_path)
    return str(output_path)


def _build_default_doc(doc: Document, reservation) -> None:
    doc.add_heading(f"Ficha de Hóspede – {settings.hotel_name}", 0)

    fields = {
        "Nome Completo": reservation.guest_name,
        "E-mail": reservation.guest_email,
        "Telefone": reservation.guest_phone,
        "Plataforma": reservation.platform,
        "ID Reserva": reservation.booking_id,
        "Check-in": reservation.checkin,
        "Check-out": reservation.checkout,
        "Quarto": reservation.room,
        "Data de Emissão": datetime.now().strftime("%d/%m/%Y %H:%M"),
    }

    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"

    for i, (label, value) in enumerate(fields.items()):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph("Assinatura do hóspede: ___________________________")
    doc.add_paragraph(f"Data: ___/___/______")


def _fill_template(doc: Document, reservation) -> None:
    replacements = {
        "{{guest_name}}": reservation.guest_name,
        "{{guest_email}}": reservation.guest_email,
        "{{guest_phone}}": reservation.guest_phone,
        "{{checkin}}": reservation.checkin,
        "{{checkout}}": reservation.checkout,
        "{{room}}": reservation.room,
        "{{booking_id}}": reservation.booking_id,
        "{{platform}}": reservation.platform,
        "{{hotel_name}}": settings.hotel_name,
    }

    def _replace_in_paragraph(para) -> None:
        for key, val in replacements.items():
            if key in para.text:
                for run in para.runs:
                    run.text = run.text.replace(key, val)

    for para in doc.paragraphs:
        _replace_in_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para)
